"""
LlamaIndex document processing pipeline.

Key improvements over raw chunking:
1. HierarchicalNodeParser — stores both summary and detail chunks
2. SentenceWindowNodeParser — preserves sentence context around each chunk
3. Section classification — tags each node with document role
4. Sparse vector generation — enables hybrid BM25 + dense search in Qdrant
"""
import re
import uuid
import hashlib
from typing import TYPE_CHECKING
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import (
    HierarchicalNodeParser,
    SentenceWindowNodeParser,
    get_leaf_nodes,
)
import httpx
from openai import OpenAI, AzureOpenAI
from app.config import settings

if TYPE_CHECKING:
    from app.core.output_models import EvaluationSetup

_embed_client = None


def get_embed_client():
    global _embed_client
    if _embed_client is None:
        http_client = httpx.Client(verify=settings.ssl_verify) if not settings.ssl_verify else None
        if settings.llm_provider == "azure":
            _embed_client = AzureOpenAI(
                azure_endpoint=settings.azure_openai_endpoint,
                api_key=settings.azure_openai_api_key,
                api_version=settings.azure_openai_api_version,
                **({"http_client": http_client} if http_client else {}),
            )
        else:
            _embed_client = OpenAI(
                api_key=settings.openai_api_key,
                **({"http_client": http_client} if http_client else {}),
            )
    return _embed_client


def get_dense_embedding(text: str) -> list[float]:
    """Dense embedding — 3072 dimensions.
    Uses Azure OpenAI when LLM_PROVIDER=azure, otherwise OpenAI directly.
    Returns zero vector when skip_embeddings=True (dev/test mode)."""
    if settings.skip_embeddings:
        return [0.0] * 3072
    client = get_embed_client()
    model = (
        settings.azure_openai_embedding_deployment
        if settings.llm_provider == "azure"
        else settings.openai_embedding_model
    )
    response = client.embeddings.create(
        model=model,
        input=text[:8000],
    )
    return response.data[0].embedding


def get_sparse_embedding(text: str) -> tuple[list[int], list[float]]:
    """
    BM25-style sparse embedding for keyword search.
    Returns (indices, values) for Qdrant sparse vector storage.
    Uses TF-IDF approximation — replace with SPLADE model for production if needed.
    """
    words = re.sub(r'[^\w\s]', ' ', text.lower()).split()
    word_freq: dict[int, float] = {}

    for word in words:
        if len(word) < 3:
            continue
        idx = int(hashlib.md5(word.encode()).hexdigest()[:8], 16) % 100000
        word_freq[idx] = word_freq.get(idx, 0) + 1.0

    if word_freq:
        max_val = max(word_freq.values())
        word_freq = {k: v / max_val for k, v in word_freq.items()}

    return list(word_freq.keys()), list(word_freq.values())


def classify_section(
    section_text: str,
    section_title: str,
    evaluation_setup: "EvaluationSetup"
) -> str:
    """
    Classifies a section as requirement_response, supporting_evidence,
    background, or boilerplate.

    Uses the customer-confirmed EvaluationSetup to identify relevant sections —
    not a hardcoded config dict. Classification adapts to whatever criteria
    the customer defined on Page 4b.
    """
    title_lower = section_title.lower()
    text_lower = section_text.lower()

    boilerplate_markers = [
        "terms and conditions", "legal notice", "disclaimer",
        "copyright", "all rights reserved", "confidentiality",
        "this document is confidential"
    ]
    if any(m in text_lower for m in boilerplate_markers):
        return "boilerplate"

    background_markers = [
        "company history", "about us", "our story", "founded in",
        "team bios", "management team", "our offices"
    ]
    if any(m in text_lower for m in background_markers):
        return "background"

    # Build keyword list from customer-defined criteria.
    # Same code, different config — logistics gets "fleet"/"delivery",
    # HR gets "payroll"/"onboarding", procurement gets "iso"/"insurance".
    all_keywords: list[str] = []

    for check in evaluation_setup.mandatory_checks:
        words = (check.name + " " + check.description).lower().split()
        all_keywords.extend(w for w in words if len(w) > 3)

    for criterion in evaluation_setup.scoring_criteria:
        words = criterion.name.lower().split()
        all_keywords.extend(w for w in words if len(w) > 3)

    for target in evaluation_setup.extraction_targets:
        words = (target.name + " " + target.description).lower().split()
        all_keywords.extend(w for w in words if len(w) > 3)

    all_keywords = list(set(all_keywords))

    for kw in all_keywords:
        if kw in title_lower or kw in text_lower:
            return "requirement_response"

    evidence_markers = [
        "certificate", "certification", "insurance", "sla",
        "service level", "case study", "project reference",
        "client testimonial", "award"
    ]
    if any(m in title_lower or m in text_lower for m in evidence_markers):
        return "supporting_evidence"

    return "background"


def process_document(
    content: bytes,
    filename: str,
    vendor_id: str,
    org_id: str,
    evaluation_setup: "EvaluationSetup"
) -> list[dict]:
    """
    Full LlamaIndex processing pipeline.

    Returns list of chunk dicts ready for Qdrant insertion:
    {
        chunk_id, text, dense_vector, sparse_indices, sparse_values,
        section_id, section_title, section_type, priority,
        page_number, filename, vendor_id, org_id
    }
    """
    raw_text = _extract_text(content, filename)

    if not raw_text or len(raw_text.strip()) < 100:
        return []

    doc = LlamaDocument(
        text=raw_text,
        metadata={
            "filename": filename,
            "vendor_id": vendor_id,
            "org_id": org_id,
        }
    )

    hierarchical_parser = HierarchicalNodeParser.from_defaults(
        chunk_sizes=[2048, 512, 128]
    )
    sentence_parser = SentenceWindowNodeParser.from_defaults(
        window_size=3,
        window_metadata_key="window",
        original_text_metadata_key="original_text"
    )

    all_nodes = hierarchical_parser.get_nodes_from_documents([doc])
    leaf_nodes = get_leaf_nodes(all_nodes)
    sentence_nodes = sentence_parser.get_nodes_from_documents([doc])

    chunks = []
    seen_texts: set[str] = set()

    for node in leaf_nodes + sentence_nodes:
        text = node.get_content().strip()

        if len(text) < 80:
            continue

        text_hash = hashlib.md5(text.encode()).hexdigest()
        if text_hash in seen_texts:
            continue
        seen_texts.add(text_hash)

        section_title = node.metadata.get(
            "section_title",
            _detect_section_title(text)
        )
        section_id = node.metadata.get(
            "section_id",
            _generate_section_id(section_title)
        )
        section_type = classify_section(text, section_title, evaluation_setup)

        priority_map = {
            "requirement_response": 1,
            "supporting_evidence": 2,
            "background": 3,
            "boilerplate": 4
        }
        priority = priority_map.get(section_type, 3)

        dense = get_dense_embedding(text)
        sparse_indices, sparse_values = get_sparse_embedding(text)

        chunks.append({
            "chunk_id": str(uuid.uuid4()),
            "text": text,
            "dense_vector": dense,
            "sparse_indices": sparse_indices,
            "sparse_values": sparse_values,
            "section_id": section_id,
            "section_title": section_title,
            "section_type": section_type,
            "priority": priority,
            "page_number": node.metadata.get("page_label", 1),
            "filename": filename,
            "vendor_id": vendor_id,
            "org_id": org_id,
            "window": node.metadata.get("window", ""),
        })

    return chunks


def _extract_text(content: bytes, filename: str) -> str:
    """Extract text from PDF, DOCX, or TXT."""
    import io
    if filename.lower().endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join(
            page.extract_text() or ""
            for page in reader.pages
        )
    elif filename.lower().endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text)
    else:
        return content.decode("utf-8", errors="ignore")


def _detect_section_title(text: str) -> str:
    """Extract section title from the beginning of a chunk."""
    lines = text.strip().split("\n")
    first_line = lines[0].strip() if lines else ""
    if len(first_line) < 100 and first_line:
        return first_line
    return "General"


def _generate_section_id(title: str) -> str:
    """Generate a short section ID from title."""
    clean = re.sub(r'[^\w\s]', '', title.lower())
    words = clean.split()[:3]
    return "-".join(words) if words else "general"
